# -*- coding: utf-8 -*-
"""
Калькулятор ОСЭЭК — Интегральный индекс социально-экономической 
эффективности коммуникаций

Авторы методики: Алтухов А.С., Бобылева А.З.
"""

import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import io
import requests

# Для экспорта в Word
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Настройка страницы
st.set_page_config(
    page_title="Калькулятор ОСЭЭК",
    page_icon="📊",
    layout="wide"
)

# ============================================================================
# КОНФИГУРАЦИЯ
# ============================================================================

# API-ключ берется из Streamlit Secrets (безопасно)
# Если ключ не настроен — функция автозаполнения по ИНН будет недоступна
DADATA_API_KEY = st.secrets.get("DADATA_API_KEY", None)

# ============================================================================
# СПРАВОЧНЫЕ ДАННЫЕ
# ============================================================================

# Отрасли с X_ref и K_risk (ОТКАЛИБРОВАННЫЕ ЗНАЧЕНИЯ)
INDUSTRY_DATA = {
    "Энергетика и ТЭК (нефть, газ, электроэнергетика)": {
        "x_ref": 60000,
        "k_risk": 1.10,
        "category": "high",
        "examples": "Газпром, Роснефть, Лукойл, Россети",
        "okved_prefixes": ["06", "09.1", "35"]
    },
    "Атомная энергетика": {
        "x_ref": 60000,
        "k_risk": 1.10,
        "category": "high",
        "examples": "Росатом",
        "okved_prefixes": ["24.46"]
    },
    "Банки и финансовые услуги": {
        "x_ref": 55000,
        "k_risk": 1.0,
        "category": "high",
        "examples": "Сбер, ВТБ, Альфа-Банк",
        "okved_prefixes": ["64", "65", "66"]
    },
    "Оборонно-промышленный комплекс": {
        "x_ref": 45000,
        "k_risk": 1.10,
        "category": "high",
        "examples": "Ростех, ОАК, ОСК",
        "okved_prefixes": ["25.4", "30.11", "30.3"]
    },
    "Телекоммуникации и IT": {
        "x_ref": 45000,
        "k_risk": 1.0,
        "category": "medium",
        "examples": "МТС, Мегафон, Ростелеком, Яндекс",
        "okved_prefixes": ["61", "62", "63"]
    },
    "Ритейл и электронная коммерция": {
        "x_ref": 40000,
        "k_risk": 1.0,
        "category": "medium",
        "examples": "X5 Group, Магнит, Ozon, Wildberries",
        "okved_prefixes": ["47"]
    },
    "Строительство и девелопмент": {
        "x_ref": 35000,
        "k_risk": 1.0,
        "category": "medium",
        "examples": "ПИК, Самолет, ЛСР, Эталон",
        "okved_prefixes": ["41", "42", "43"]
    },
    "Металлургия и горнодобыча": {
        "x_ref": 35000,
        "k_risk": 1.10,
        "category": "medium",
        "examples": "Норникель, НЛМК, Северсталь, АЛРОСА",
        "okved_prefixes": ["07", "08", "24"]
    },
    "Транспорт и логистика": {
        "x_ref": 30000,
        "k_risk": 1.0,
        "category": "medium",
        "examples": "РЖД, Аэрофлот, FESCO",
        "okved_prefixes": ["49", "50", "51", "52"]
    },
    "Химическая промышленность": {
        "x_ref": 25000,
        "k_risk": 1.10,
        "category": "medium",
        "examples": "Сибур, ФосАгро, Уралхим",
        "okved_prefixes": ["20", "21"]
    },
    "Машиностроение": {
        "x_ref": 25000,
        "k_risk": 1.0,
        "category": "medium",
        "examples": "КАМАЗ, ГАЗ, Трансмашхолдинг",
        "okved_prefixes": ["28", "29"]
    },
    "Фармацевтика и медицина": {
        "x_ref": 25000,
        "k_risk": 1.0,
        "category": "medium",
        "examples": "Р-Фарм, Биокад, Герофарм",
        "okved_prefixes": ["21", "86"]
    },
    "Агропромышленный комплекс": {
        "x_ref": 20000,
        "k_risk": 1.0,
        "category": "low",
        "examples": "Русагро, Черкизово, Мираторг",
        "okved_prefixes": ["01", "10", "11"]
    },
    "Табачная промышленность": {
        "x_ref": 15000,
        "k_risk": 1.10,
        "category": "low",
        "examples": "",
        "okved_prefixes": ["12"]
    },
    "Алкогольная промышленность": {
        "x_ref": 15000,
        "k_risk": 1.10,
        "category": "low",
        "examples": "",
        "okved_prefixes": ["11.0"]
    },
    "Другая отрасль": {
        "x_ref": 20000,
        "k_risk": 1.0,
        "category": "low",
        "examples": "",
        "okved_prefixes": []
    }
}

# Критерии транспарентности
TRANSP_CRITERIA = [
    "Годовой отчёт о деятельности компании опубликован на официальном сайте",
    "Финансовая отчётность заверена внешним аудитором",
    "Опубликован отчёт об устойчивом развитии (ESG/КСО)",
    "Отчётность соответствует международным стандартам (GRI, SASB)",
    "Раскрыта структура собственности",
    "Раскрыт состав органов управления (с биографиями или без)",
    "Раскрыто вознаграждение топ-менеджмента",
    "Раздел существенных фактов (события, влияющие на стоимость акций) обновляется регулярно",
    "На сайте указаны контакты для инвесторов, СМИ, соискателей",
    "Доступна англоязычная версия годового отчёта или сайта"
]

# Критерии институциональной зрелости (ОБНОВЛЕННЫЙ СПИСОК)
INST_CRITERIA = [
    ("В структуре компании есть подразделение по коммуникациям", 10),
    ("Руководитель по коммуникациям входит в состав топ-менеджмента", 10),
    ("На сайте компании публикуются корпоративные новости и указаны контакты пресс-службы", 10),
    ("На сайте компании регулярно публикуются корпоративные новости — не реже 4 раз в месяц", 10),
    ("Компания ведет официальное сообщество в VK", 5),
    ("Компания ведет официальный Telegram-канал", 5),
    ("Коммуникационная стратегия компании публично доступна (на сайте или в годовом отчете)", 10),
    ("Предусмотрен механизм обратной связи: горячая линия, контакт-центр, форма обращений и др.", 10),
    ("Существуют антикризисные коммуникационные процедуры (для внутренней оценки)", 10),
    ("Руководитель по коммуникациям или компания состоит в профессиональной ассоциации (РАСО, АКМР, АКОС и др.)", 10),
    ("Награды или признание в области коммуникаций за последние 3 года", 10),
    ("KPI коммуникационного подразделения привязаны к бизнес-показателям компании (для внутренней оценки)", 10),
    ("Проводится регулярный (не реже раза в год) мониторинг репутации или восприятия стейкхолдерами", 10),
]

# ============================================================================
# ФУНКЦИИ ДЛЯ РАБОТЫ С DaData API
# ============================================================================

def get_company_by_inn(inn: str) -> dict:
    """Получение данных о компании по ИНН через DaData API"""
    
    # Проверка наличия API-ключа
    if not DADATA_API_KEY:
        return {"error": "API-ключ DaData не настроен. Функция автозаполнения недоступна."}
    
    if not inn.isdigit() or len(inn) not in [10, 12]:
        return {"error": "Неверный формат ИНН. Должно быть 10 цифр (юрлицо) или 12 цифр (ИП)"}
    
    try:
        url = "https://suggestions.dadata.ru/suggestions/api/4_1/rs/findById/party"
        headers = {
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Authorization": f"Token {DADATA_API_KEY}"
        }
        data = {"query": inn}
        
        response = requests.post(url, json=data, headers=headers, timeout=10)
        
        if response.status_code == 200:
            result = response.json()
            if result.get("suggestions") and len(result["suggestions"]) > 0:
                company = result["suggestions"][0]
                company_data = company.get("data", {})
                
                return {
                    "name": company.get("value", ""),
                    "full_name": company_data.get("name", {}).get("full_with_opf", ""),
                    "okved": company_data.get("okved", ""),
                    "okved_name": company_data.get("okved_type", ""),
                    "address": company_data.get("address", {}).get("value", ""),
                    "status": company_data.get("state", {}).get("status", ""),
                    "employees": company_data.get("employee_count", None)
                }
            else:
                return {"error": "Компания с таким ИНН не найдена"}
        elif response.status_code == 401:
            return {"error": "Ошибка авторизации API. Проверьте ключ."}
        elif response.status_code == 403:
            return {"error": "Доступ запрещен. Проверьте лимиты API."}
        else:
            return {"error": f"Ошибка сервера: {response.status_code}"}
            
    except requests.exceptions.Timeout:
        return {"error": "Превышено время ожидания ответа от сервера"}
    except requests.exceptions.RequestException as e:
        return {"error": f"Ошибка соединения: {str(e)}"}
    except Exception as e:
        return {"error": f"Неизвестная ошибка: {str(e)}"}

def determine_industry_by_okved(okved: str) -> str:
    """Определение отрасли по коду ОКВЭД"""
    if not okved:
        return "Другая отрасль"
    
    for industry, data in INDUSTRY_DATA.items():
        for prefix in data.get("okved_prefixes", []):
            if okved.startswith(prefix):
                return industry
    return "Другая отрасль"

# ============================================================================
# ФУНКЦИИ РАСЧЕТА ОСЭЭК
# ============================================================================

def calculate_i_media(val_i: float, x_ref: float) -> float:
    if x_ref <= 0:
        return 0.0
    i_media = (val_i / x_ref) * 100
    if i_media < 0:
        i_media = 0.0
    elif i_media > 100:
        i_media = 100.0
    return i_media

def calculate_v_vol(monthly_values: list, x_ref: float = None) -> float:
    """
    Расчет коэффициента волатильности.
    Если среднее значение < 1% от X_ref (или < 1 при отсутствии X_ref),
    возвращается максимальная волатильность (данные недостаточны для анализа).
    """
    if len(monthly_values) < 2:
        return 0.0
    mu = np.mean(monthly_values)
    
    # Порог: 1% от эталона или минимум 1
    threshold = (x_ref * 0.01) if x_ref and x_ref > 0 else 1.0
    
    if mu < threshold:
        return float('inf')  # Данные недостаточны для корректного расчета
    
    sigma = np.std(monthly_values, ddof=0)
    v_vol = sigma / mu
    return v_vol

def calculate_m_stab(i_media: float, v_vol: float) -> float:
    if v_vol == float('inf'):
        return 0.0
    m_stab = i_media / (1 + v_vol)
    return m_stab

def calculate_v_hr(rank: int, total: int) -> float:
    if total <= 1:
        return 0.0
    v_hr = (1 - (rank - 1) / (total - 1)) * 100
    return max(0.0, min(100.0, v_hr))

def calculate_r_transp(indicators: list) -> float:
    return sum([10 if ind else 0 for ind in indicators])

def calculate_r_inst(indicators: list, scores: list) -> float:
    total = 0
    for ind, score in zip(indicators, scores):
        if ind:
            total += score
    return min(total, 100)  # Ограничиваем 100 баллами

def calculate_s_rep(v_hr: float, r_transp: float, r_inst: float) -> float:
    return (v_hr + r_transp + r_inst) / 3

def calculate_i_core(m_stab: float, s_rep: float) -> float:
    return m_stab * 0.6 + s_rep * 0.4

def calculate_i_adj(i_core: float, k_risk: float, k_scale: float) -> float:
    return i_core * k_risk * k_scale

def get_k_scale(employees: int, is_strategic: bool) -> float:
    """Получение коэффициента масштаба (ОТКАЛИБРОВАННОЕ ЗНАЧЕНИЕ)"""
    if employees >= 100000 or is_strategic:
        return 1.05
    return 1.0

def calculate_roi(revenue: float, costs: float) -> float:
    if costs <= 0:
        return 0.0
    return ((revenue - costs) / costs) * 100

def calculate_sroi(social_value: float, costs: float) -> float:
    if costs <= 0:
        return 0.0
    return ((social_value - costs) / costs) * 100

def get_k_roi(roi: float) -> float:
    return 0.05 if roi > 0 else 0.0

def get_k_sroi(sroi: float) -> float:
    return 0.10 if sroi > 0 else 0.0

def get_k_budget(plan: float, fact: float, has_approval: bool) -> float:
    if plan <= 0:
        return 0.0
    deviation = ((fact - plan) / plan) * 100
    if deviation > 10 and not has_approval:
        return -0.10
    return 0.0

def calculate_k_eff(k_roi: float, k_sroi: float, k_budget: float) -> float:
    return 1 + k_roi + k_sroi + k_budget

def calculate_iseec_e(iseec_b: float, k_eff: float) -> float:
    return iseec_b * k_eff

def get_quality_rating(value: float) -> tuple:
    """Возвращает (уровень, эмодзи, цвет)"""
    if value > 100:
        return "Очень высокий", "🟢", "#28a745"
    elif value >= 76:
        return "Высокий", "🟢", "#28a745"
    elif value >= 51:
        return "Средний", "🟡", "#ffc107"
    elif value >= 26:
        return "Низкий", "🟠", "#fd7e14"
    else:
        return "Критически низкий", "🔴", "#dc3545"

def calculate_manual_track_i_media(n_pos: int, n_neg: int, n_total: int) -> float:
    if n_total <= 0:
        return 0.0
    i_media = (1 + (n_pos - n_neg) / n_total) * 50
    return max(0.0, min(100.0, i_media))

def generate_recommendations(transp_indicators: list, inst_indicators: list, inst_scores: list, v_hr: float, m_stab: float) -> list:
    """Генерация топ-5 рекомендаций на основе невыполненных критериев"""
    recommendations = []
    
    for i, (done, label) in enumerate(zip(transp_indicators, TRANSP_CRITERIA)):
        if not done:
            recommendations.append({
                "area": "Транспарентность",
                "action": label,
                "potential": 10,
                "priority": 2
            })
    
    for i, ((label, score), done) in enumerate(zip(INST_CRITERIA, inst_indicators)):
        if not done:
            recommendations.append({
                "area": "Институциональная зрелость",
                "action": label,
                "potential": score,
                "priority": 2 if score == 10 else 3
            })
    
    if v_hr < 50:
        recommendations.append({
            "area": "HR-бренд",
            "action": "Принять участие в Рейтинге работодателей России (hh.ru)",
            "potential": 50,
            "priority": 1
        })
    
    if m_stab < 50:
        recommendations.append({
            "area": "Медийная активность",
            "action": "Усилить присутствие в СМИ, увеличить количество публикаций",
            "potential": 30,
            "priority": 1
        })
    
    recommendations.sort(key=lambda x: (x["priority"], -x["potential"]))
    
    return recommendations[:5]

def generate_word_report(r: dict, recommendations: list) -> bytes:
    """Генерация расширенного отчета в формате Word"""
    doc = Document()
    
    title = doc.add_heading('ОТЧЕТ О РАСЧЕТЕ ОСЭЭК', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph()
    
    doc.add_heading('1. Информация о компании', level=1)
    
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'
    
    data1 = [
        ('Название компании', r['company_name'] or 'Не указано'),
        ('Отчетный период', f"{r['report_year']} год"),
        ('Отрасль', r['industry']),
        ('Численность сотрудников', f"{r['employees']:,}"),
        ('Стратегическое предприятие', 'Да' if r['is_strategic'] else 'Нет')
    ]
    
    for i, (label, value) in enumerate(data1):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    doc.add_heading('2. Итоговые показатели', level=1)
    
    p = doc.add_paragraph()
    p.add_run(f"ISEEC_B (базовый контур): ").bold = True
    p.add_run(f"{r['iseec_b']:.1f} баллов — {r['rating_b']}")
    
    if r['iseec_e']:
        p = doc.add_paragraph()
        p.add_run(f"ISEEC_E (расширенный контур): ").bold = True
        p.add_run(f"{r['iseec_e']:.1f} баллов — {r['rating_e']}")
    
    doc.add_paragraph()
    
    doc.add_heading('Шкала интерпретации', level=2)
    
    scale_table = doc.add_table(rows=6, cols=3)
    scale_table.style = 'Table Grid'
    
    scale_data = [
        ('Баллы', 'Уровень', 'Интерпретация'),
        ('> 100', 'Очень высокий', 'Коммуникации создают дополнительную ценность'),
        ('76–100', 'Высокий', 'Эффективная коммуникационная система'),
        ('51–75', 'Средний', 'Есть резервы для улучшения'),
        ('26–50', 'Низкий', 'Требуется существенная доработка'),
        ('0–25', 'Критически низкий', 'Коммуникационная система неэффективна'),
    ]
    
    for i, row_data in enumerate(scale_data):
        for j, cell_text in enumerate(row_data):
            scale_table.rows[i].cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('3. Детализация расчета', level=1)
    
    doc.add_heading('3.1. Субиндекс медийной устойчивости (M_stab)', level=2)
    doc.add_paragraph(f"I_media: {r['i_media']:.1f}")
    doc.add_paragraph(f"V_vol: {r['v_vol']:.3f}")
    doc.add_paragraph(f"M_stab: {r['m_stab']:.1f}")
    
    doc.add_heading('3.2. Субиндекс социальной репутации (S_rep)', level=2)
    doc.add_paragraph(f"V_hr: {r['v_hr']:.1f}")
    doc.add_paragraph(f"R_transp: {r['r_transp']:.0f}")
    doc.add_paragraph(f"R_inst: {r['r_inst']:.0f}")
    doc.add_paragraph(f"S_rep: {r['s_rep']:.1f}")
    
    doc.add_heading('3.3. Формула расчета', level=2)
    doc.add_paragraph(f"I_Core = M_stab × 0.6 + S_rep × 0.4 = {r['i_core']:.1f}")
    doc.add_paragraph(f"K_risk = {r['k_risk']:.2f}, K_scale = {r['k_scale']:.2f}")
    doc.add_paragraph(f"ISEEC_B = I_Core × K_risk × K_scale = {r['iseec_b']:.1f}")
    
    if r['iseec_e']:
        doc.add_heading('3.4. Расширенный контур', level=2)
        doc.add_paragraph(f"K_eff: {r['k_eff']:.2f}")
        doc.add_paragraph(f"ISEEC_E = ISEEC_B × K_eff = {r['iseec_e']:.1f}")
    
    doc.add_paragraph()
    
    doc.add_heading('4. Потенциал роста', level=1)
    
    potential_m = max(0, 100 - r['m_stab'])
    potential_transp = max(0, 100 - r['r_transp'])
    potential_inst = max(0, 100 - r['r_inst'])
    potential_hr = max(0, 100 - r['v_hr'])
    
    doc.add_paragraph(f"Медийная устойчивость (M_stab): {r['m_stab']:.0f}/100 — потенциал +{potential_m:.0f} баллов")
    doc.add_paragraph(f"Транспарентность (R_transp): {r['r_transp']:.0f}/100 — потенциал +{potential_transp:.0f} баллов")
    doc.add_paragraph(f"Институциональная зрелость (R_inst): {r['r_inst']:.0f}/100 — потенциал +{potential_inst:.0f} баллов")
    doc.add_paragraph(f"HR-бренд (V_hr): {r['v_hr']:.0f}/100 — потенциал +{potential_hr:.0f} баллов")
    
    doc.add_paragraph()
    
    if recommendations:
        doc.add_heading('5. Рекомендации по улучшению', level=1)
        
        for i, rec in enumerate(recommendations[:5], 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. [{rec['area']}] ").bold = True
            p.add_run(f"{rec['action']}")
            p.add_run(f" (+{rec['potential']} баллов)")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('Методика ОСЭЭК').italic = True
    doc.add_paragraph('Авторы: Алтухов А.С., Бобылева А.З.')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ============================================================================
# СТРАНИЦЫ ПРИЛОЖЕНИЯ
# ============================================================================

def page_calculator():
    """Основная страница калькулятора"""
    
    st.title("📊 Калькулятор ОСЭЭК")
    st.markdown("""
    **Интегральный индекс социально-экономической эффективности коммуникаций**
    
    ---
    **Авторы методики:** Алтухов А.С., Бобылева А.З.
    """)
    
    st.divider()
    
    # ========================================================================
    # БЛОК 1: ИНФОРМАЦИЯ О КОМПАНИИ
    # ========================================================================
    st.header("1️⃣ Информация о компании")
    
    if 'auto_company_name' not in st.session_state:
        st.session_state.auto_company_name = ""
    if 'auto_industry' not in st.session_state:
        st.session_state.auto_industry = None
    if 'auto_employees' not in st.session_state:
        st.session_state.auto_employees = None
    
    with st.expander("🔍 Автозаполнение по ИНН (опционально)"):
        st.markdown("""
        Введите ИНН компании для автоматического получения названия и определения отрасли.
        Если ИНН не указан — заполните данные вручную ниже.
        """)
        
        col_inn1, col_inn2 = st.columns([2, 1])
        with col_inn1:
            inn_input = st.text_input(
                "ИНН компании",
                placeholder="10 или 12 цифр",
                max_chars=12,
                help="Введите ИНН для автоматического получения данных из ЕГРЮЛ"
            )
        with col_inn2:
            st.write("")
            st.write("")
            inn_button = st.button("Найти компанию", use_container_width=True)
        
        if inn_button and inn_input:
            with st.spinner("Поиск данных в ЕГРЮЛ..."):
                result = get_company_by_inn(inn_input)
                if "error" in result:
                    st.error(result["error"])
                else:
                    st.success(f"✅ Найдена компания: **{result.get('name', 'Н/Д')}**")
                    st.session_state.auto_company_name = result.get('name', '')
                    
                    okved = result.get('okved', '')
                    if okved:
                        detected_industry = determine_industry_by_okved(okved)
                        st.session_state.auto_industry = detected_industry
                        st.info(f"📌 ОКВЭД: {okved} → Отрасль: {detected_industry}")
                    
                    if result.get('employees'):
                        st.session_state.auto_employees = result.get('employees')
                        st.info(f"👥 Численность сотрудников: {result.get('employees')}")
                    
                    if result.get('address'):
                        st.caption(f"Адрес: {result.get('address')}")
    
    col1, col2 = st.columns(2)
    with col1:
        default_name = st.session_state.auto_company_name if st.session_state.auto_company_name else ""
        company_name = st.text_input(
            "Название компании",
            value=default_name,
            placeholder="Введите название организации",
            help="Укажите название в любом формате. Поле необязательное — используется для отчета."
        )
    with col2:
        report_year = st.selectbox(
            "Отчетный период",
            options=list(range(2025, 2019, -1)),
            index=0
        )
    
    st.subheader("Параметры организации")
    
    col1, col2 = st.columns(2)
    with col1:
        industry_list = list(INDUSTRY_DATA.keys())
        default_industry_index = 0
        if st.session_state.auto_industry and st.session_state.auto_industry in industry_list:
            default_industry_index = industry_list.index(st.session_state.auto_industry)
        
        industry = st.selectbox(
            "Отрасль деятельности",
            options=industry_list,
            index=default_industry_index,
            help="Выбор отрасли определяет коэффициент K_risk и эталонное значение X_ref"
        )
        
        ind_data = INDUSTRY_DATA[industry]
        if ind_data["examples"]:
            st.caption(f"Примеры: {ind_data['examples']}")
    
    with col2:
        default_employees = st.session_state.auto_employees if st.session_state.auto_employees else 10000
        employees = st.number_input(
            "Численность сотрудников",
            min_value=0,
            value=default_employees,
            step=1000,
            help="Организации с численностью ≥ 100 000 получают K_scale = 1.05"
        )
    
    is_strategic = st.checkbox(
        "Стратегическое предприятие (включено в перечень Указа Президента РФ)",
        help="Стратегические предприятия получают K_scale = 1.05 независимо от численности"
    )
    
    k_risk = ind_data["k_risk"]
    k_scale = get_k_scale(employees, is_strategic)
    
    with st.expander("📌 Коэффициенты для вашей организации", expanded=True):
        st.markdown(f"""
**K_risk = {k_risk}** — коэффициент отраслевого риска.  
Учитывает повышенное общественное внимание к социально значимым отраслям (атомная энергетика, ОПК, химия и др.). Значение 1,10 означает +10% к оценке за работу в сложных условиях.

**K_scale = {k_scale}** — коэффициент масштаба.  
Учитывает сложность координации коммуникаций в крупных организациях (≥100 000 сотрудников) или стратегических предприятиях. Значение 1,05 означает +5% к оценке.

**Оба коэффициента применяются одновременно** (перемножаются): итоговый множитель = {k_risk} × {k_scale} = {k_risk * k_scale:.3f}
        """)
    
    st.divider()
    
    # ========================================================================
    # БЛОК 2: МЕДИАДАННЫЕ
    # ========================================================================
    st.header("2️⃣ Медиаданные (для расчета M_stab)")
    
    with st.expander("ℹ️ Что такое Медиалогия?"):
        st.markdown("""
        **[Медиалогия](https://www.mlg.ru)** — ведущая российская система мониторинга и анализа СМИ. 
        
        **МедиаИндекс** — ключевой показатель системы, который оценивает качество присутствия 
        компании в СМИ по трем параметрам:
        - Индекс цитируемости источника
        - Тональность упоминания (позитив/негатив)
        - Заметность сообщения
        
        **Как получить данные:**
        1. Войдите в систему Медиалогия
        2. Выберите объект мониторинга (вашу компанию)
        3. Перейдите в раздел «Аналитика»
        4. Выберите период (год) и выгрузите отчет
        
        Если нет доступа к Медиалогии — используйте режим «Ручной ввод».
        """)
    
    media_source = st.radio(
        "Источник данных:",
        options=["У меня есть данные из Медиалогии", "Ручной ввод (Manual Track)"]
    )
    
    default_x_ref = ind_data["x_ref"]
    
    if media_source == "У меня есть данные из Медиалогии":
        col1, col2 = st.columns(2)
        with col1:
            media_index_year = st.number_input(
                "МедиаИндекс за год (Val_i)",
                min_value=-1000000.0,
                value=45000.0,
                step=1000.0,
                help="Суммарный МедиаИндекс компании за отчетный год."
            )
        
        with col2:
            use_custom_xref = st.checkbox("Ввести X_ref вручную")
            
            if use_custom_xref:
                x_ref = st.number_input(
                    "Эталон отрасли (X_ref)",
                    min_value=1.0,
                    value=float(default_x_ref),
                    step=1000.0
                )
            else:
                x_ref = float(default_x_ref)
                st.info(f"X_ref = {default_x_ref:,} (ориентировочное значение для отрасли)")
        
        with st.expander("📖 Как рассчитать X_ref самостоятельно?"):
            st.markdown(f"""
            **X_ref** — эталонное значение МедиаИндекса для сравнения.
            
            **Алгоритм:**
            1. Определите 3 крупнейших компании вашей отрасли
            2. Найдите их пиковые годовые МедиаИндексы за 3 года
            3. Рассчитайте среднее арифметическое
            
            **Сейчас используется:** X_ref = {x_ref:,.0f}
            """)
        
        st.markdown("**Помесячные значения МедиаИндекса** *(для расчёта волатильности V_vol)*:")
        st.caption("Волатильность показывает стабильность медиаприсутствия. Резкие скачки снижают итоговую оценку.")
        
        months = ["Янв", "Фев", "Мар", "Апр", "Май", "Июн", "Июл", "Авг", "Сен", "Окт", "Ноя", "Дек"]
        cols = st.columns(6)
        monthly_values = []
        
        for i, month in enumerate(months):
            with cols[i % 6]:
                val = st.number_input(month, min_value=-100000.0, value=3900.0, step=100.0, key=f"month_{i}")
                monthly_values.append(val)
        
        i_media = calculate_i_media(media_index_year, x_ref)
        v_vol = calculate_v_vol(monthly_values, x_ref)
        m_stab = calculate_m_stab(i_media, v_vol)
        
    else:
        st.info("📌 **Ручной ввод** — соберите данные о публикациях через новостные агрегаторы (Google News и др.) или поиск по СМИ.")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            n_total = st.number_input("Всего публикаций за год", min_value=0, value=100)
        with col2:
            n_pos = st.number_input("Позитивных и нейтральных", min_value=0, value=70)
        with col3:
            n_neg = st.number_input("Негативных", min_value=0, value=30)
        
        st.markdown("**Количество публикаций по месяцам** *(для расчёта волатильности)*:")
        st.caption("Волатильность показывает стабильность медиаприсутствия. Резкие скачки снижают итоговую оценку.")
        
        months = ["Янв", "Фев", "Мар", "Апр", "Май", "Июн", "Июл", "Авг", "Сен", "Окт", "Ноя", "Дек"]
        cols = st.columns(6)
        monthly_values = []
        
        for i, month in enumerate(months):
            with cols[i % 6]:
                val = st.number_input(month, min_value=0, value=8, step=1, key=f"manual_month_{i}")
                monthly_values.append(val)
        
        i_media = calculate_manual_track_i_media(n_pos, n_neg, n_total)
        v_vol = calculate_v_vol(monthly_values)
        m_stab = calculate_m_stab(i_media, v_vol)
        x_ref = None
    
    st.divider()
    
    # ========================================================================
    # БЛОК 3: HR-БРЕНД
    # ========================================================================
    st.header("3️⃣ HR-бренд (V_hr)")
    
    st.markdown("**[Рейтинг работодателей России](https://rating.hh.ru)** — ежегодное исследование HeadHunter.")
    
    hr_source = st.radio(
        "Статус компании в рейтинге:",
        options=[
            "Компания есть в Рейтинге работодателей России (hh.ru)",
            "Компании нет в рейтинге — крупная или публичная организация",
            "Компании нет в рейтинге — средняя или непубличная организация",
            "Другой случай / не знаю"
        ]
    )
    
    if hr_source == "Компания есть в Рейтинге работодателей России (hh.ru)":
        col1, col2 = st.columns(2)
        with col1:
            hr_rank = st.number_input("Место в рейтинге", min_value=1, value=45)
        with col2:
            hr_total = st.number_input("Всего участников", min_value=1, value=700)
        v_hr = calculate_v_hr(hr_rank, hr_total)
    elif "крупная или публичная" in hr_source:
        v_hr = 0.0
        st.info("ℹ️ Для крупных публичных организаций отсутствие в рейтинге — сигнал о проблемах. V_hr = 0")
    elif "средняя или непубличная" in hr_source:
        v_hr = 50.0
        st.info("ℹ️ Для средних и непубличных организаций: V_hr = 50 (нейтрально)")
    else:
        v_hr = st.slider("Укажите V_hr вручную:", min_value=0.0, max_value=100.0, value=50.0, step=5.0)
    
    st.divider()
    
    # ========================================================================
    # БЛОК 4: ТРАНСПАРЕНТНОСТЬ
    # ========================================================================
    st.header("4️⃣ Транспарентность (R_transp)")
    st.markdown("Отметьте выполненные критерии. Каждый пункт = 10 баллов. **Если ничего не выполняется — оставьте пустым.**")
    
    transp_indicators = []
    cols = st.columns(2)
    for i, label in enumerate(TRANSP_CRITERIA):
        with cols[i % 2]:
            val = st.checkbox(label, key=f"transp_{i}")
            transp_indicators.append(val)
    
    r_transp = calculate_r_transp(transp_indicators)
    
    st.divider()
    
    # ========================================================================
    # БЛОК 5: ИНСТИТУЦИОНАЛЬНАЯ ЗРЕЛОСТЬ
    # ========================================================================
    st.header("5️⃣ Институциональная зрелость (R_inst)")
    st.markdown("Отметьте выполненные критерии. **Максимум 100 баллов.** Если ничего не выполняется — оставьте пустым.")
    
    inst_indicators = []
    inst_scores = []
    
    cols = st.columns(2)
    for i, (label, score) in enumerate(INST_CRITERIA):
        with cols[i % 2]:
            val = st.checkbox(label, key=f"inst_{i}")
            inst_indicators.append(val)
            inst_scores.append(score)
    
    r_inst = calculate_r_inst(inst_indicators, inst_scores)
    
    st.divider()
    
    # ========================================================================
    # БЛОК 6: РАСШИРЕННЫЙ КОНТУР
    # ========================================================================
    st.header("6️⃣ Расширенный контур (опционально)")
    
    st.info("""
    ℹ️ **Расширенный контур** требует внутренних данных: ROI, SROI, бюджет.
    
    - Оцениваете **свою компанию** — включите раздел
    - Анализируете **чужую компанию** — пропустите
    
    **Укажите данные за выбранный в начале отчетный период.**
    """)
    
    use_extended = st.checkbox("Рассчитать расширенный контур ISEEC_E")
    
    k_roi_val = 0.0
    k_sroi_val = 0.0
    k_budget_val = 0.0
    
    if use_extended:
        st.subheader("ROI коммуникаций")
        roi_method = st.radio("Способ ввода ROI:", ["Ввести ROI (%)", "Рассчитать", "Нет данных"], key="roi_method")
        
        if roi_method == "Ввести ROI (%)":
            roi_value = st.number_input("ROI (%)", value=15.0)
            k_roi_val = get_k_roi(roi_value)
        elif roi_method == "Рассчитать":
            col1, col2 = st.columns(2)
            with col1:
                revenue = st.number_input("Доходы от коммуникаций (млн руб.)", value=520.0, min_value=0.0)
            with col2:
                costs = st.number_input("Затраты (млн руб.)", value=450.0, min_value=0.0)
            roi_value = calculate_roi(revenue, costs)
            st.success(f"ROI: {roi_value:.1f}%")
            k_roi_val = get_k_roi(roi_value)
        
        st.subheader("SROI коммуникаций")
        sroi_method = st.radio("Способ ввода SROI:", ["Ввести SROI (%)", "Рассчитать", "Нет данных"], key="sroi_method")
        
        if sroi_method == "Ввести SROI (%)":
            sroi_value = st.number_input("SROI (%)", value=40.0)
            k_sroi_val = get_k_sroi(sroi_value)
        elif sroi_method == "Рассчитать":
            col1, col2 = st.columns(2)
            with col1:
                social_value = st.number_input("Социальная ценность (млн руб.)", value=630.0, min_value=0.0)
            with col2:
                costs_sroi = st.number_input("Затраты (млн руб.)", value=450.0, min_value=0.0, key="costs_sroi")
            sroi_value = calculate_sroi(social_value, costs_sroi)
            st.success(f"SROI: {sroi_value:.1f}%")
            k_sroi_val = get_k_sroi(sroi_value)
        
        st.subheader("Бюджетная дисциплина")
        col1, col2 = st.columns(2)
        with col1:
            budget_plan = st.number_input("План (млн руб.)", value=430.0, min_value=0.0)
        with col2:
            budget_fact = st.number_input("Факт (млн руб.)", value=450.0, min_value=0.0)
        
        has_approval = st.checkbox("Превышение согласовано руководством")
        k_budget_val = get_k_budget(budget_plan, budget_fact, has_approval)
    
    st.divider()
    
    # ========================================================================
    # РАСЧЕТ И РЕЗУЛЬТАТЫ
    # ========================================================================
    if st.button("🧮 Рассчитать ОСЭЭК", type="primary", use_container_width=True):
        
        s_rep = calculate_s_rep(v_hr, r_transp, r_inst)
        i_core = calculate_i_core(m_stab, s_rep)
        i_adj = calculate_i_adj(i_core, k_risk, k_scale)
        iseec_b = i_adj
        
        if use_extended:
            k_eff = calculate_k_eff(k_roi_val, k_sroi_val, k_budget_val)
            iseec_e = calculate_iseec_e(iseec_b, k_eff)
        else:
            k_eff = None
            iseec_e = None
        
        rating_b, emoji_b, color_b = get_quality_rating(iseec_b)
        
        st.header("📊 Результаты расчета")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric(label="ISEEC_B (базовый контур)", value=f"{iseec_b:.1f}")
            st.markdown(f"**{emoji_b} {rating_b}**")
        
        with col2:
            if iseec_e is not None:
                rating_e, emoji_e, color_e = get_quality_rating(iseec_e)
                st.metric(label="ISEEC_E (расширенный)", value=f"{iseec_e:.1f}")
                st.markdown(f"**{emoji_e} {rating_e}**")
            else:
                st.metric(label="ISEEC_E", value="—")
                st.caption("Не рассчитывался")
        
        with col3:
            st.metric(label="I_Core (ядро)", value=f"{i_core:.1f}")
            st.caption(f"K_risk={k_risk}, K_scale={k_scale}")
        
        # Визуальная шкала
        st.subheader("📏 Позиция на шкале")
        
        # Определяем максимум шкалы
        scale_max = 125 if iseec_b > 100 else 100
        marker_pos = min(iseec_b / scale_max * 100, 100)
        
        if scale_max == 125:
            scale_html = f"""
            <div style="margin: 20px 0;">
                <div style="display: flex; justify-content: space-between; font-size: 12px; color: #666; margin-bottom: 5px;">
                    <span>0</span>
                    <span>25</span>
                    <span>50</span>
                    <span>75</span>
                    <span>100</span>
                    <span>125</span>
                </div>
                <div style="position: relative; height: 30px; background: linear-gradient(to right, #dc3545 0%, #fd7e14 20%, #ffc107 40%, #28a745 60%, #28a745 80%, #1a5c1a 100%); border-radius: 5px;">
                    <div style="position: absolute; left: {marker_pos}%; top: -5px; transform: translateX(-50%);">
                        <div style="width: 0; height: 0; border-left: 8px solid transparent; border-right: 8px solid transparent; border-top: 10px solid #333;"></div>
                    </div>
                    <div style="position: absolute; left: {marker_pos}%; top: 35px; transform: translateX(-50%); font-weight: bold; font-size: 14px;">
                        {iseec_b:.1f}
                    </div>
                </div>
                <div style="display: flex; justify-content: space-between; font-size: 10px; color: #999; margin-top: 25px;">
                    <span>Критический</span>
                    <span>Низкий</span>
                    <span>Средний</span>
                    <span>Высокий</span>
                    <span>Очень высокий</span>
                    <span>Превосходный</span>
                </div>
            </div>
            """
        else:
            scale_html = f"""
            <div style="margin: 20px 0;">
                <div style="display: flex; justify-content: space-between; font-size: 12px; color: #666; margin-bottom: 5px;">
                    <span>0</span>
                    <span>25</span>
                    <span>50</span>
                    <span>75</span>
                    <span>100</span>
                </div>
                <div style="position: relative; height: 30px; background: linear-gradient(to right, #dc3545 0%, #fd7e14 25%, #ffc107 50%, #28a745 75%, #28a745 100%); border-radius: 5px;">
                    <div style="position: absolute; left: {marker_pos}%; top: -5px; transform: translateX(-50%);">
                        <div style="width: 0; height: 0; border-left: 8px solid transparent; border-right: 8px solid transparent; border-top: 10px solid #333;"></div>
                    </div>
                    <div style="position: absolute; left: {marker_pos}%; top: 35px; transform: translateX(-50%); font-weight: bold; font-size: 14px;">
                        {iseec_b:.1f}
                    </div>
                </div>
                <div style="display: flex; justify-content: space-between; font-size: 10px; color: #999; margin-top: 25px;">
                    <span>Критический</span>
                    <span>Низкий</span>
                    <span>Средний</span>
                    <span>Высокий</span>
                    <span>Очень высокий</span>
                </div>
            </div>
            """
        st.markdown(scale_html, unsafe_allow_html=True)
        
        # Детализация
        st.subheader("📋 Детализация по компонентам")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Медийная устойчивость (M_stab)**")
            st.progress(min(m_stab / 100, 1.0))
            st.caption(f"{m_stab:.1f} / 100 баллов")
            
            st.markdown("**Транспарентность (R_transp)**")
            st.progress(min(r_transp / 100, 1.0))
            st.caption(f"{r_transp:.0f} / 100 баллов")
        
        with col2:
            st.markdown("**HR-бренд (V_hr)**")
            st.progress(min(v_hr / 100, 1.0))
            st.caption(f"{v_hr:.1f} / 100 баллов")
            
            st.markdown("**Институциональная зрелость (R_inst)**")
            st.progress(min(r_inst / 100, 1.0))
            st.caption(f"{r_inst:.0f} / 100 баллов")
        
        # Потенциал роста
        st.subheader("📈 Потенциал роста")
        
        potential_data = [
            ("Медийная устойчивость", m_stab, 100 - m_stab),
            ("Транспарентность", r_transp, 100 - r_transp),
            ("Институциональная зрелость", r_inst, 100 - r_inst),
            ("HR-бренд", v_hr, 100 - v_hr),
        ]
        
        potential_data.sort(key=lambda x: x[2], reverse=True)
        
        for name, current, potential in potential_data:
            if potential > 0:
                col1, col2, col3 = st.columns([3, 1, 1])
                with col1:
                    st.write(name)
                with col2:
                    st.write(f"{current:.0f} / 100")
                with col3:
                    st.write(f"**+{potential:.0f}** возможно")
        
        # Профиль коммуникаций
        st.subheader("🎯 Профиль коммуникаций")
        
        diff_ms_sr = m_stab - s_rep
        
        if abs(diff_ms_sr) <= 15:
            st.info("✅ **Сбалансированный профиль.** Медийная активность соответствует репутационному капиталу. Это оптимальное состояние.")
        elif diff_ms_sr > 15:
            st.info("📢 **Медийное доминирование.** Активность в СМИ опережает институциональную базу. Рекомендуется усилить работу над транспарентностью и институциональной зрелостью.")
        else:
            st.info("🏛️ **Репутационное доминирование.** Институциональная база сильнее медийного присутствия. Рекомендуется усилить работу со СМИ и медиаактивность.")
        
        # Рекомендации
        recommendations = generate_recommendations(transp_indicators, inst_indicators, inst_scores, v_hr, m_stab)
        
        if recommendations:
            st.subheader("💡 Рекомендации по улучшению")
            
            for i, rec in enumerate(recommendations[:5], 1):
                with st.container():
                    col1, col2 = st.columns([5, 1])
                    with col1:
                        st.markdown(f"**{i}. [{rec['area']}]** {rec['action']}")
                    with col2:
                        st.markdown(f"**+{rec['potential']}** б.")
        
        # Шкала интерпретации
        with st.expander("📖 Шкала интерпретации ОСЭЭК"):
            st.markdown("""
            | Баллы | Уровень | Что это значит |
            |-------|---------|----------------|
            | **> 100** | Очень высокий | Коммуникации создают дополнительную ценность для компании |
            | **76–100** | Высокий | Эффективная коммуникационная система, соответствующая лучшим практикам |
            | **51–75** | Средний | Система работает, но есть значительные резервы для улучшения |
            | **26–50** | Низкий | Требуется существенная доработка коммуникационной системы |
            | **0–25** | Критически низкий | Коммуникационная система неэффективна, нужны срочные меры |
            """)
        
        # Сохранение результатов
        st.session_state['results'] = {
            'company_name': company_name,
            'report_year': report_year,
            'industry': industry,
            'employees': employees,
            'is_strategic': is_strategic,
            'm_stab': m_stab,
            'i_media': i_media,
            'v_vol': v_vol,
            'v_hr': v_hr,
            'r_transp': r_transp,
            'r_inst': r_inst,
            's_rep': s_rep,
            'i_core': i_core,
            'k_risk': k_risk,
            'k_scale': k_scale,
            'i_adj': i_adj,
            'iseec_b': iseec_b,
            'iseec_e': iseec_e,
            'k_eff': k_eff,
            'rating_b': rating_b,
            'rating_e': get_quality_rating(iseec_e)[0] if iseec_e else None
        }
        st.session_state['recommendations'] = recommendations
    
    # ========================================================================
    # ЭКСПОРТ
    # ========================================================================
    if 'results' in st.session_state:
        st.divider()
        st.header("📄 Экспорт отчета")
        
        r = st.session_state['results']
        recs = st.session_state.get('recommendations', [])
        
        safe_name = (r['company_name'] or 'company').replace(' ', '_').replace('"', '').replace('«', '').replace('»', '')[:30]
        
        report_text = f"""ОТЧЕТ О РАСЧЕТЕ ОСЭЭК
{'='*60}

Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}

1. ИНФОРМАЦИЯ О КОМПАНИИ
Название: {r['company_name'] or 'Не указано'}
Период: {r['report_year']} год
Отрасль: {r['industry']}
Сотрудников: {r['employees']:,}

2. РЕЗУЛЬТАТЫ
ISEEC_B: {r['iseec_b']:.1f} баллов — {r['rating_b']}
{'ISEEC_E: ' + f"{r['iseec_e']:.1f} баллов — {r['rating_e']}" if r['iseec_e'] else ''}

3. ШКАЛА ИНТЕРПРЕТАЦИИ
> 100    — Очень высокий (коммуникации создают ценность)
76–100   — Высокий (эффективная система)
51–75    — Средний (есть резервы)
26–50    — Низкий (требуется доработка)
0–25     — Критически низкий (неэффективно)

4. ДЕТАЛИЗАЦИЯ
M_stab (медийная устойчивость): {r['m_stab']:.1f}/100
V_hr (HR-бренд): {r['v_hr']:.1f}/100
R_transp (транспарентность): {r['r_transp']:.0f}/100
R_inst (институциональная зрелость): {r['r_inst']:.0f}/100

5. ТОП-5 РЕКОМЕНДАЦИЙ
"""
        for i, rec in enumerate(recs[:5], 1):
            report_text += f"{i}. [{rec['area']}] {rec['action']} (+{rec['potential']} б.)\n"
        
        report_text += f"""
{'='*60}
Методика ОСЭЭК
Авторы: Алтухов А.С., Бобылева А.З.
"""
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if DOCX_AVAILABLE:
                word_bytes = generate_word_report(r, recs)
                st.download_button(
                    label="📥 Скачать (Word)",
                    data=word_bytes,
                    file_name=f"OSEEK_{safe_name}_{r['report_year']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        with col2:
            st.download_button(
                label="📥 Скачать (TXT)",
                data=report_text,
                file_name=f"OSEEK_{safe_name}_{r['report_year']}.txt",
                mime="text/plain"
            )
        
        with col3:
            csv_data = pd.DataFrame([r]).to_csv(index=False)
            st.download_button(
                label="📥 Скачать (CSV)",
                data=csv_data,
                file_name=f"OSEEK_{safe_name}_{r['report_year']}.csv",
                mime="text/csv"
            )


def page_methodology():
    """Страница методики"""
    
    st.title("📖 О методике ОСЭЭК")
    
    st.markdown("""
    ## Интегральный индекс социально-экономической эффективности коммуникаций
    
    **ОСЭЭК** (ISEEC — Integral Index of Socio-Economic Effectiveness of Communications) — 
    методика количественной оценки результативности коммуникационной системы организации.
    
    ### Назначение
    
    Методика позволяет оценить, насколько эффективно коммуникации организации выполняют 
    три ключевые институциональные функции:
    
    1. **Снижение информационной асимметрии** между организацией и стейкхолдерами
    2. **Легитимация экономических решений** в глазах общества
    3. **Согласование интересов** организации и ключевых социальных групп
    
    ### Двухконтурная архитектура
    
    | Контур | Данные | Применение |
    |--------|--------|------------|
    | **ISEEC_B** (базовый) | Публичные источники | Внешняя оценка, бенчмаркинг |
    | **ISEEC_E** (расширенный) | + управленческий учет | Внутренняя оценка, KPI |
    
    ### Структура индекса
    
    ```
    ISEEC_B = I_Core × K_risk × K_scale
    
    где I_Core = M_stab × 0.6 + S_rep × 0.4
    ```
    
    **Коэффициенты корректировки:**
    - K_risk = 1.10 для социально значимых отраслей (энергетика, ОПК, химия и др.), 1.0 для остальных
    - K_scale = 1.05 для крупных (≥100 000 сотрудников) и стратегических предприятий, 1.0 для остальных
    
    ### Шкала интерпретации
    
    | Баллы | Уровень | Интерпретация |
    |-------|---------|---------------|
    | > 100 | Очень высокий | Коммуникации создают дополнительную ценность |
    | 76–100 | Высокий | Эффективная коммуникационная система |
    | 51–75 | Средний | Есть резервы для улучшения |
    | 26–50 | Низкий | Требуется существенная доработка |
    | 0–25 | Критически низкий | Коммуникационная система неэффективна |
    
    ---
    
    ### Авторы методики
    
    **Алтухов А.С., Бобылева А.З.**
    """)


# ============================================================================
# ГЛАВНАЯ ФУНКЦИЯ
# ============================================================================

def main():
    page = st.sidebar.radio(
        "Навигация",
        options=["🧮 Калькулятор", "📖 О методике ОСЭЭК"],
        index=0
    )
    
    st.sidebar.divider()
    st.sidebar.markdown("""
    **Калькулятор ОСЭЭК**  
    Версия 1.1
    
    ---
    Авторы методики:  
    Алтухов А.С., Бобылева А.З.
    """)
    
    if page == "🧮 Калькулятор":
        page_calculator()
    else:
        page_methodology()


if __name__ == "__main__":
    main()

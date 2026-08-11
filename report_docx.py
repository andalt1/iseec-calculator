# -*- coding: utf-8 -*-
"""Выгрузка расчетного протокола ОСЭЭК в формате Word."""

import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from oseec_core import (LEVEL_NAMES, MONTHS, TRANSP_BLOCK1, TRANSP_BLOCK2,
                        INST_BLOCK1, INST_BLOCK2, fmt)

NAVY = RGBColor(0x01, 0x21, 0x69)
INK = RGBColor(0x23, 0x28, 0x30)


def _style(doc: Document) -> None:
    st_ = doc.styles["Normal"]
    st_.font.name = "Times New Roman"
    st_.font.size = Pt(12)
    st_.font.color.rgb = INK
    for lvl, size in (("Heading 1", 16), ("Heading 2", 13)):
        h = doc.styles[lvl]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.color.rgb = NAVY
        h.font.bold = True


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11.5)
        for r in cells[1].paragraphs[0].runs:
            r.font.bold = True


def build_report(res: dict, company: str, base) -> bytes:
    doc = Document()
    _style(doc)

    h = doc.add_heading("Протокол расчета индекса ОСЭЭК", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Интегральный индекс социально-экономической эффективности "
        "коммуникаций. Методика: Алтухов А.С., Бобылева А.З. "
        "Свидетельство о государственной регистрации программы для ЭВМ "
        "№ 2026663079 от 04.05.2026."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(10.5)

    doc.add_heading("1. Объект оценки", level=2)
    _kv_table(doc, [
        ("Организация", company or "не указана"),
        ("Дата расчета", date.today().strftime("%d.%m.%Y")),
        ("Контур оценки", "Расширенный (базовый + управленческий учет)"
         if "oseec_e" in res else "Базовый (открытые источники)"),
    ])

    doc.add_heading("2. Компоненты и субиндексы", level=2)
    rows = [
        ("Индекс медийной результативности I_media",
         fmt(res["i_media"]) if res.get("i_media") is not None else "не рассчитывался"),
        ("Коэффициент волатильности V_vol",
         fmt(res["v_vol"]) if res.get("v_vol") is not None else "не рассчитывался"),
        ("Субиндекс медийной устойчивости M_stab", fmt(res["m_stab"])),
        ("Верификация HR-бренда V_hr", fmt(res["v_hr"]) +
         (" (сценарное значение)" if res.get("hr_scenario_mode") else "")),
        ("Транспарентность R_transp", fmt(res["r_transp"])),
        ("Институциональная зрелость R_inst", fmt(res["r_inst"])),
        ("Субиндекс социальной репутации S_rep", fmt(res["s_rep"])),
    ]
    _kv_table(doc, rows)

    doc.add_heading("3. Расчет итогового значения", level=2)
    calc_rows = [
        ("Ядро индекса I_Core = 0,6 × M_stab + 0,4 × S_rep", fmt(res["i_core"])),
        ("Коэффициент риска коммуникационной среды K_risk", fmt(res["k_risk"])),
        ("Коэффициент масштаба организации K_scale", fmt(res["k_scale"])),
        ("Базовый контур ОСЭЭК_B = I_Core × K_risk × K_scale",
         fmt(res["oseec_b"])),
        ("Качественный уровень (базовый контур)", res["level_name"]),
    ]
    if "oseec_e" in res:
        calc_rows += [
            ("Коэффициент управленческой эффективности K_eff", fmt(res["k_eff"])),
            ("Расширенный контур ОСЭЭК_E = ОСЭЭК_B × K_eff", fmt(res["oseec_e"])),
            ("Качественный уровень (расширенный контур)", res["level_e_name"]),
        ]
    _kv_table(doc, calc_rows)

    doc.add_heading("4. Диагностические статусы", level=2)
    if res["critical"]:
        for kind, name, val in res["critical"]:
            doc.add_paragraph(
                f"Критическое ограничение: {kind} «{name}» — значение "
                f"{fmt(val)} балла ниже порогового уровня "
                f"({'40' if kind == 'субиндекс' else '30'} баллов).",
                style="List Bullet")
    else:
        doc.add_paragraph(
            "Значения субиндексов и компонентов выше пороговых уровней, "
            "статусы критических ограничений не присвоены.")

    if res.get("hr_scenarios"):
        doc.add_heading("5. Сценарный диапазон компонента V_hr", level=2)
        doc.add_paragraph(
            "Компания не представлена в рейтингах работодателей как "
            "самостоятельное юридическое лицо, поэтому по компоненту "
            "верификации HR-бренда применен сценарный подход раздела 3.2 "
            "методики.")
        _kv_table(doc, [
            ("ОСЭЭК_B при V_hr = 0", fmt(res["hr_scenarios"][0.0])),
            ("ОСЭЭК_B при V_hr = 50 (базовый сценарий)",
             fmt(res["hr_scenarios"][50.0])),
            ("ОСЭЭК_B при V_hr = 100", fmt(res["hr_scenarios"][100.0])),
        ])

    doc.add_heading("Приложение. Исходные данные", level=2)
    if base.media_track == "manual":
        doc.add_paragraph(
            "Медиакорпус (ручной протокол), помесячное распределение "
            "«всего / в том числе негативных»:")
        t = doc.add_table(rows=2, cols=13)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Месяц"
        t.rows[1].cells[0].text = "Публикации"
        for i in range(12):
            t.rows[0].cells[i + 1].text = MONTHS[i][:3]
            t.rows[1].cells[i + 1].text = (f"{base.monthly_total[i]}/"
                                           f"{base.monthly_neg[i]}")
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                        r.font.name = "Times New Roman"
    elif base.media_track == "monitoring":
        doc.add_paragraph(
            f"Данные системы медиамониторинга: X_fact = {fmt(base.x_fact)}, "
            f"X_ref = {fmt(base.x_ref)}.")
    else:
        doc.add_paragraph(
            "Публикации, атрибутированные к юридическому лицу, отсутствуют — "
            "применено правило нулевого значения субиндекса медийной "
            "устойчивости.")

    def _mark(items, values):
        return "; ".join(f"{name} — {'да' if v else 'нет'}"
                         for name, v in zip(items, values))

    doc.add_paragraph("Чек-лист транспарентности, блок корпоративной "
                      "открытости: " + _mark(TRANSP_BLOCK1, base.transp_b1) + ".")
    doc.add_paragraph("Чек-лист транспарентности, блок компенсаторного "
                      "раскрытия: " + _mark(TRANSP_BLOCK2, base.transp_b2) + ".")
    doc.add_paragraph("Чек-лист институциональной зрелости, блок каналов и "
                      "процедур обратной связи: "
                      + _mark(INST_BLOCK1, base.inst_b1) + ".")
    doc.add_paragraph("Чек-лист институциональной зрелости, блок "
                      "институционального закрепления: "
                      + _mark(INST_BLOCK2, base.inst_b2) + ".")

    note = doc.add_paragraph(
        "Расчет выполнен онлайн-калькулятором ОСЭЭК по формулам (1)–(11) "
        "методики. Значения компонентов зафиксированы с точностью до сотых "
        "долей, производные величины рассчитаны без промежуточного "
        "округления.")
    for r in note.runs:
        r.font.size = Pt(10)
        r.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
